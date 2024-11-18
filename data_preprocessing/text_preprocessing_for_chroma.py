import io
import PyPDF2
from PyPDF2 import PdfReader
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
import logging

# Initialize logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger()


class FileProcessor:
    """
    A class to handle different file types and convert them to plain text.
    """
    @staticmethod
    def process_file(file):
        """
        Process the uploaded file based on its extension (PDF, DOCX, TXT).
        """
        try:
            file_name = file.filename.lower()

            # Process the file according to its extension
            if file_name.endswith('.pdf'):
                return await PdfFileProcessor().convert_to_text(file)
            elif file_name.endswith('.docx'):
                return await DocxFileProcessor().convert_to_text(file)
            elif file_name.endswith('.txt'):
                return await TxtFileProcessor().process(file)
            else:
                logger.error(f"Unsupported file type: {file_name}")
                return None
        except Exception as e:
            logger.error(f"Error processing file: {e}")
            return None


class DocxFileProcessor:
    """
    A class to process DOCX files and convert them to plain text.
    """
    @staticmethod
    async def convert_to_text(docx_file):
        """
        Convert DOCX file to plain text.
        """
        try:
            doc = Document(io.BytesIO(await docx_file.read()))
            text = ""
            # Extract text from paragraphs and tables
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + '\n'
            return text
        except Exception as e:
            logger.error(f"Error processing DOCX file: {e}")
            return None


class PdfFileProcessor:
    """
    A class to process PDF files and convert them to plain text.
    """
    @staticmethod
    async def convert_to_text(pdf_file):
        """
        Convert PDF file to plain text.
        """
        try:
            pdf_file.seek(0)
            pdf_reader = PdfReader(pdf_file)
            text = ""
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text()
            return text
        except Exception as e:
            logger.error(f"Error converting PDF to text: {e}")
            return None


class TxtFileProcessor:
    """
    A class to process TXT files and convert them to plain text.
    """
    @staticmethod
    async def process(txt_file):
        """
        Process TXT file and return its contents as a string.
        """
        try:
            # Read the file content asynchronously
            file_content = await txt_file.read()

            # Convert the bytes to a string using utf-8 encoding
            text = file_content.decode("utf-8", errors="replace")
            return text
        except Exception as e:
            logger.error(f"Error processing TXT file: {e}")
            return None


class ChunkHandler:
    """
    A class to handle chunking of text using the Langchain text splitter.
    """
    @staticmethod
    def langchain_create_chunk_from_text(all_text, chunk_size=1000, chunk_overlap=100):
        """
        Split text into chunks using Langchain's RecursiveCharacterTextSplitter.

        :param all_text: The full text to be split into chunks.
        :param chunk_size: The size of each chunk (default 1000 characters).
        :param chunk_overlap: The number of characters to overlap between chunks (default 100).
        :return: A list of text chunks.
        """
        try:
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                length_function=len
            )
            logger.info(f"Splitting text into chunks with size={chunk_size} and overlap={chunk_overlap}")
            chunk_list = text_splitter.create_documents([all_text])
            return chunk_list
        except Exception as e:
            logger.error(f"Error creating chunks: {e}")
            return None
